from flask import Flask, render_template, request, redirect, url_for, session, jsonify
# from werkzeug.security import generate_password_hash, check_password_hash
# from werkzeug.utils import secure_filename
import os
import openai
from openai import OpenAI
# from PyPDF2 import PdfReaders
import pandas as pd
import ollama
import json
from collections import OrderedDict
from difflib import SequenceMatcher

app = Flask(__name__)
# app.secret_key = 'your_secret_key'
model = 'gpt-4o-mini'

# Configure upload folder
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# User database (in-memory for simplicity)
users_db = {}

# OpenAI API key setup
openai.api_key = "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"

# Routes
@app.route('/')
def home():
    return render_template('dashboard.html')

@app.route('/dashboard', methods=['GET', 'POST'])
def dashboard():

    if request.method == 'POST':
        # File upload logic
        if 'file' not in request.files:
            return "No file uploaded!"
        file = request.files['file']
        if file.filename == '':
            return "No selected file!"
        if file:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            # Analyze the file
            regulations_path = './local/list_regulations.docx'
            weights_csv = './local/Adjusted_Importance-Based_Regulation_Weights.csv'
            total_score, completeness_score, transparency_score, analysis_result, detailed_analysis = analyze_file(filepath, regulations_path, weights_csv)
            return render_template('results.html', result=analysis_result, detailed_analysis=detailed_analysis, total_score=total_score, completeness_score=completeness_score, transparency_score=transparency_score
)

    return render_template('dashboard.html')


@app.route('/chat', methods=['POST'])
def chat():

    data = request.get_json()
    user_input = data.get("message", "")

    if not user_input.strip():
        return jsonify({"reply": "Empty message."})

    client = OpenAI(api_key=openai.api_key)

    response = client.chat.completions.create(
        model= model,
        messages=[
            {"role": "system", "content": "You are a helpful assistant for legal and contract-related questions."},
            {"role": "user", "content": user_input}
        ],
        max_tokens=500
    )

    reply = response.choices[0].message.content
    return jsonify({"reply": reply})


@app.route('/extract_options', methods=['POST'])
def extract_options():

    data = request.get_json() or {}
    question = (data.get("question") or "").strip()
    num_options = data.get("num_options", 0)

    try:
        num_options = int(num_options)
    except (TypeError, ValueError):
        num_options = 0

    if not question:
        return jsonify({"error": "Scenario text is empty."}), 400
    if num_options <= 0:
        return jsonify({"error": "Number of options must be a positive integer."}), 400

    client = OpenAI(api_key=openai.api_key)

    # Ask the LLM to extract EXACT options from the question text
    response = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are an assistant that STRICTLY extracts multiple-choice answer options "
                    "from a survey question. You MUST not paraphrase or alter the wording, "
                    "capitalization, or punctuation of any option text."
                )
            },
            {
                "role": "user",
                "content": (
                    "Here is a survey question. Extract exactly "
                    f"{num_options} distinct answer option texts from it. "
                    "The options you return MUST be exactly the same substrings as in the original "
                    "question text (no rephrasing, no editing). "
                    "Return ONLY a compact JSON array of strings with length "
                    f"{num_options} and no other text.\n\n"
                    f"Question:\n{question}"
                )
            },
        ],
        max_tokens=300,
        temperature=0,
    )

    raw = response.choices[0].message.content.strip()

    try:
        options = json.loads(raw)
        if not isinstance(options, list):
            raise ValueError("Not a list")
    except Exception:
        # If parsing fails, return a clean error so the UI can tell the user
        return jsonify({
            "error": "Failed to parse options from model response.",
            "raw_response": raw
        }), 500

    # Trim each option, just whitespace only – content stays identical
    options = [o.strip() for o in options]

    return jsonify({"options": options})


@app.route('/simulate', methods=['POST'])
def simulate():

    data = request.get_json() or {}
    prompt = (data.get("prompt") or "").strip()

    # options from frontend (option boxes)
    raw_options = data.get("options") or []
    if not isinstance(raw_options, list):
        raw_options = []

    options = [str(o).strip() for o in raw_options if str(o).strip()]
    num_options = len(options)

    # fallback: if no options sent, use numeric count
    if num_options == 0:
        try:
            num_options = int(data.get("num_options", 0))
        except (TypeError, ValueError):
            num_options = 0

    if not prompt:
        return jsonify({"error": "Empty prompt."}), 400
    if num_options <= 0:
        return jsonify({"error": "Number of options must be a positive integer."}), 400

    # clamp simulations
    try:
        num_simulations = int(data.get("num_simulations", 20))
    except (TypeError, ValueError):
        num_simulations = 20
    num_simulations = max(1, min(num_simulations, 100))

    client = OpenAI(api_key=openai.api_key)

    # ---- 1) Get N answers from LLM as *indices* 1..num_options ----

    answers_raw = []

    # Build an options block if we have explicit texts
    if options:
        options_block = "\n".join(
            f"{i+1}. {opt}" for i, opt in enumerate(options)
        )
    else:
        options_block = ""  # question may already contain them
    for _ in range(num_simulations):
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a survey participant. Think as human with varity"
                        f"This is a multiple-choice question with {num_options} options. "
                        "Choose exactly one option. "
                        "Answer ONLY with the chosen option label and text as format 'label: text' "
                        "The answer text must be exactly the same as original option text."
                         "The label must be a simple alphanumeric label such as '1', '2', 'A', 'B'. "
                        "If the question does not explicitly provide labels, create numerical labels for options "
                        "in order (1, 2, 3, ...). "
                        "Do NOT add explanations, reasoning, sentences, or anything else. "
                        "Output exactly one line, for example: '1: Yes'."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            max_tokens=50,
            temperature=1.0,  # randomness so we get variation
        )
        ans = response.choices[0].message.content.strip()
        print(ans)
        answers_raw.append(ans)

    # ---- 2) Convert answers to indices and count per option ----

    def extract_option_text(answer: str) -> str:
        """
        From a model answer like '1: some text' or 'A: some text',
        return just 'some text'. If there is no ':', return the whole string.
        """
        s = (answer or "").strip()
        if ":" in s:
            return s.split(":", 1)[1].strip()
        return s

    def similarity(a: str, b: str) -> float:
        """
        Case-insensitive similarity score between two strings in [0, 1].
        """
        return SequenceMatcher(None, a.lower(), b.lower()).ratio()

    counts = [0] * num_options
    other_count = 0
    SIM_THRESHOLD = 0.95  # 95% similarity

    for raw in answers_raw:
        ans_text = extract_option_text(raw)
        if not ans_text:
            other_count += 1
            continue

        best_idx = None
        best_sim = 0.0

        for i, opt in enumerate(options):
            sim = similarity(ans_text, opt)
            if sim > best_sim:
                best_sim = sim
                best_idx = i

        if best_idx is not None and best_sim >= SIM_THRESHOLD:
            counts[best_idx] += 1
        else:
            other_count += 1

    # ---- 3) Build distribution with labels "i: option text" ----
    distribution = OrderedDict()
    for i in range(num_options):
        text = options[i] if i < len(options) else ""
        label = f"{i+1}: {text}" if text else f"{i+1}"
        distribution[label] = counts[i]

    return jsonify({
        "distribution": distribution,
        "answers": answers_raw,
        "options": options,
        "num_simulations": num_simulations,
        "num_options": num_options,
    })
    

from docx import Document

def parse_regulations(file_path):
    # Extract regulations from the Word document
    doc = Document(file_path)
    regulations = {}
    current_section = None

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        # Identify headings (e.g., 'Heading 1', 'Heading 2')
        if paragraph.style.name.startswith("Heading"):
            current_section = text
            regulations[current_section] = []
        elif current_section:
            # Add content under the current section (keep bullets if present)
            if paragraph.style.name in ["List Bullet", "List Number"]:
                regulations[current_section].append({"type": "bullet", "text": text})
            else:
                regulations[current_section].append({"type": "text", "text": text})

    return regulations



def analyze_file(filepath, regulations_path, weights_csv):
    # Read the file content
    ext = os.path.splitext(filepath)[1]
    if ext == '.txt':
        with open(filepath, 'r', encoding="utf8") as file:
            content = file.read()
    # elif ext == '.pdf':
    #     reader = PdfReader(filepath)
    #     content = " ".join(page.extract_text() for page in reader.pages)
    elif ext in ['.doc', '.docx']:
        import docx
        doc = docx.Document(filepath)
        content = " ".join([p.text for p in doc.paragraphs])
    else:
        return "Unsupported file type!"
    
    if not content.strip():
        return "The file is empty. Please upload a valid agreement text."

    # Send content to GPT4
    client = OpenAI(api_key='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')

    # Define the messages for the chat completion
    messages = [
        {"role": "system", "content": "You are an expert legal analyst."},
        {"role": "user", "content": f"Please analyze the following agreement:\n\n{content}"}
    ]

    # Create a chat completion
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        max_tokens=500
    )

    # Extract and print the response content
    summary  = response.choices[0].message.content


    # Parse regulations
    regulations = parse_regulations(regulations_path)

    weights = pd.read_csv(weights_csv)
    weight_dict = dict(zip(weights['Regulation'], weights['Weight']))

    # Initialize variables for scoring
    completeness_score = 0
    total_weight = sum(weight_dict.values())
    detailed_analysis = {}

    # Analyze each regulation
    detailed_analysis = {}
    for section, points in regulations.items():
        section_analysis = []
        for point in points:
            if point["type"] == "text":
                reg_response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "You are an expert legal analyst."},
                        {"role": "user", "content": f"Does the following document cover this regulation? Regulation: {point['text']}\n\nDocument: {content}"}
                    ],
                    max_tokens=300
                )
                reg_analysis = reg_response.choices[0].message.content
                # Check if the regulation is covered
                if "yes" in reg_analysis or "covered" in reg_analysis:
                    regulation_name = point['text']
                    completeness_score += weight_dict.get(regulation_name, 0)
                section_analysis.append({"type": "text", "regulation": point["text"], "analysis": reg_analysis})
            elif point["type"] == "bullet":
                reg_response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "You are an expert legal analyst."},
                        {"role": "user", "content": f"Does the following document cover this regulation? Regulation: {point['text']}\n\nDocument: {content}"}
                    ],
                    max_tokens=300
                )
                reg_analysis = reg_response.choices[0].message.content
                # Check if the regulation is covered
                if "yes" in reg_analysis or "covered" in reg_analysis:
                    regulation_name = point['text']
                    completeness_score += weight_dict.get(regulation_name, 0)
                section_analysis.append({"type": "bullet", "regulation": point["text"], "analysis": reg_analysis})
        detailed_analysis[section] = section_analysis
    
    # Scale completeness score to 50 points
    completeness_score = round(completeness_score * 50 / total_weight, 2)

    # Evaluate transparency for the entire document
    transparency_response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a legal language evaluator."},
            {"role": "user", "content": f"Evaluate the transparency of the following document, the response should be just a number between 0 and 50, no more other words in response:\n\n{content}"}
        ],
        max_tokens=300
    )
    transparency_score = float(transparency_response.choices[0].message.content.split()[-1])  # Assuming it returns a score
    # transparency_score = round(min(max(transparency_score, 0), 50), 2)

    # Total Score
    total_score = completeness_score + transparency_score
        
        # Return the response content
    return total_score, completeness_score, transparency_score, summary,detailed_analysis 

if __name__ == '__main__':
    app.run(debug=True)